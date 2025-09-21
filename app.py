import streamlit as st
import pandas as pd
import sqlite3
import json
import re
import os
import logging
from datetime import datetime
from typing import Dict, List, Tuple, Optional
import hashlib
import io
import zipfile
import tempfile
from pathlib import Path

# Default API Key Configuration
DEFAULT_GEMINI_API_KEY = "AIzaSyDAx61-09OGYB0J6ab2BvgPI3ZIHM7MTYg"

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Document processing with better error handling
DEPENDENCIES_AVAILABLE = {'pdf': False, 'docx': False, 'genai': False}

try:
    import PyPDF2
    DEPENDENCIES_AVAILABLE['pdf'] = True
except ImportError:
    logger.warning("PyPDF2 not available")

try:
    from docx import Document
    DEPENDENCIES_AVAILABLE['docx'] = True
except ImportError:
    logger.warning("python-docx not available")

try:
    import google.generativeai as genai
    DEPENDENCIES_AVAILABLE['genai'] = True
except ImportError:
    logger.warning("google-generativeai not available")

# Page configuration
st.set_page_config(
    page_title="Resume Relevance Check System",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Check dependencies and show warnings
def check_dependencies():
    """Check if all required dependencies are available"""
    missing_deps = []
    
    if not DEPENDENCIES_AVAILABLE['pdf']:
        missing_deps.append("PyPDF2")
    if not DEPENDENCIES_AVAILABLE['docx']:
        missing_deps.append("python-docx")
    if not DEPENDENCIES_AVAILABLE['genai']:
        missing_deps.append("google-generativeai")
    
    if missing_deps:
        st.error(f"Missing dependencies: {', '.join(missing_deps)}")
        st.code(f"pip install {' '.join(missing_deps)}")
        return False
    return True

# Utility functions for validation
def validate_file_size(file, max_size_mb=10):
    """Validate file size"""
    if file.size > max_size_mb * 1024 * 1024:
        raise ValueError(f"File size exceeds {max_size_mb}MB limit")
    return True

def validate_text_input(text, min_length=10, max_length=50000):
    """Validate text input"""
    if not text or not text.strip():
        raise ValueError("Text cannot be empty")
    if len(text.strip()) < min_length:
        raise ValueError(f"Text must be at least {min_length} characters long")
    if len(text.strip()) > max_length:
        raise ValueError(f"Text exceeds maximum length of {max_length} characters")
    return True

def sanitize_filename(filename):
    """Sanitize filename for safe storage"""
    import re
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    return filename[:100]  # Limit filename length

# Database setup with improved error handling
@st.cache_resource
def init_database():
    """Initialize SQLite database with required tables and error handling"""
    try:
        # Create database directory if it doesn't exist
        db_path = Path('data')
        db_path.mkdir(exist_ok=True)
        
        conn = sqlite3.connect('data/resume_relevance.db', check_same_thread=False)
        cursor = conn.cursor()
        
        # Enable foreign key constraints
        cursor.execute('PRAGMA foreign_keys = ON')
        
        # Create tables with better constraints
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS job_descriptions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL CHECK(length(title) > 0),
                company TEXT DEFAULT '',
                location TEXT DEFAULT '',
                description TEXT NOT NULL CHECK(length(description) > 10),
                must_have_skills TEXT DEFAULT '[]',
                good_to_have_skills TEXT DEFAULT '[]',
                qualifications TEXT DEFAULT '[]',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS resume_evaluations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id INTEGER NOT NULL,
                candidate_name TEXT NOT NULL CHECK(length(candidate_name) > 0),
                resume_text TEXT NOT NULL,
                relevance_score INTEGER CHECK(relevance_score >= 0 AND relevance_score <= 100),
                missing_skills TEXT DEFAULT '[]',
                matching_skills TEXT DEFAULT '[]',
                verdict TEXT CHECK(verdict IN ('High', 'Medium', 'Low')),
                suggestions TEXT DEFAULT '[]',
                strengths TEXT DEFAULT '[]',
                improvement_areas TEXT DEFAULT '[]',
                file_name TEXT DEFAULT '',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (job_id) REFERENCES job_descriptions (id) ON DELETE CASCADE
            )
        ''')
        
        # Create indexes for better performance
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_job_created_at ON job_descriptions(created_at)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_eval_job_id ON resume_evaluations(job_id)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_eval_score ON resume_evaluations(relevance_score)')
        
        conn.commit()
        logger.info("Database initialized successfully")
        return conn
        
    except Exception as e:
        logger.error(f"Database initialization failed: {str(e)}")
        st.error(f"Database initialization failed: {str(e)}")
        raise

class DocumentProcessor:
    """Handle PDF and DOCX document processing with enhanced error handling"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_file) -> str:
        """Extract text from PDF file with improved error handling"""
        try:
            if not DEPENDENCIES_AVAILABLE['pdf']:
                raise ImportError("PyPDF2 not installed")
                
            validate_file_size(pdf_file, max_size_mb=10)
            
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise ValueError("PDF file contains no pages")
                
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text + "\n"
                except Exception as page_error:
                    logger.warning(f"Error reading page {page_num + 1}: {str(page_error)}")
                    continue
            
            if not text.strip():
                raise ValueError("No text could be extracted from PDF")
                
            return text.strip()
            
        except Exception as e:
            error_msg = f"Error reading PDF: {str(e)}"
            logger.error(error_msg)
            st.error(error_msg)
            return ""
    
    @staticmethod
    def extract_text_from_docx(docx_file) -> str:
        """Extract text from DOCX file with improved error handling"""
        try:
            if not DEPENDENCIES_AVAILABLE['docx']:
                raise ImportError("python-docx not installed")
                
            validate_file_size(docx_file, max_size_mb=10)
            
            doc = Document(docx_file)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise ValueError("No text could be extracted from DOCX")
                
            return text.strip()
            
        except Exception as e:
            error_msg = f"Error reading DOCX: {str(e)}"
            logger.error(error_msg)
            st.error(error_msg)
            return ""
    
    @staticmethod
    def extract_text_from_file(uploaded_file) -> str:
        """Extract text from uploaded file based on type with validation"""
        try:
            if not uploaded_file:
                raise ValueError("No file uploaded")
                
            # Validate file type
            allowed_types = {
                "application/pdf": "pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx"
            }
            
            if uploaded_file.type not in allowed_types:
                raise ValueError(f"Unsupported file format: {uploaded_file.type}. Please upload PDF or DOCX files.")
            
            file_type = allowed_types[uploaded_file.type]
            
            # Extract text based on file type
            if file_type == "pdf":
                return DocumentProcessor.extract_text_from_pdf(uploaded_file)
            elif file_type == "docx":
                return DocumentProcessor.extract_text_from_docx(uploaded_file)
            else:
                raise ValueError("Unsupported file format")
                
        except Exception as e:
            error_msg = f"File processing error: {str(e)}"
            logger.error(error_msg)
            st.error(error_msg)
            return ""

class GeminiAnalyzer:
    """Handle Gemini API interactions for resume analysis with enhanced validation"""
    
    def __init__(self, api_key: str):
        if not api_key or len(api_key.strip()) < 10:
            raise ValueError("Please provide a valid Gemini API key")
        
        if not DEPENDENCIES_AVAILABLE['genai']:
            raise ImportError("google-generativeai not installed")
        
        try:
            genai.configure(api_key=api_key.strip())
            # Try gemini-1.5-flash first, fallback to other available models
            try:
                self.model = genai.GenerativeModel('gemini-1.5-flash')
                self.model_name = 'gemini-1.5-flash'
            except Exception:
                try:
                    self.model = genai.GenerativeModel('gemini-1.5-pro')
                    self.model_name = 'gemini-1.5-pro'
                except Exception:
                    try:
                        # Try the latest flash model
                        self.model = genai.GenerativeModel('gemini-2.0-flash-exp')
                        self.model_name = 'gemini-2.0-flash-exp'
                    except Exception:
                        # List available models for debugging
                        try:
                            models = genai.list_models()
                            available_models = [m.name for m in models if 'generateContent' in m.supported_generation_methods]
                            logger.error(f"Available models: {available_models}")
                            if available_models:
                                model_name = available_models[0].split('/')[-1]  # Get model name without prefix
                                self.model = genai.GenerativeModel(model_name)
                                self.model_name = model_name
                            else:
                                raise ValueError("No compatible models available")
                        except Exception as list_error:
                            logger.error(f"Failed to list models: {list_error}")
                            raise ValueError("Could not initialize any Gemini model")
            
            self.api_key = api_key.strip()
            logger.info(f"Gemini API initialized successfully with model: {self.model_name}")
        except Exception as e:
            logger.error(f"Gemini API initialization failed: {str(e)}")
            raise ValueError(f"Failed to initialize Gemini API: {str(e)}")
    
    def get_available_models(self):
        """Get list of available models for debugging"""
        try:
            models = genai.list_models()
            return [m.name for m in models if 'generateContent' in m.supported_generation_methods]
        except Exception as e:
            logger.error(f"Error listing models: {str(e)}")
            return []
    
    def _clean_json_response(self, response_text: str) -> str:
        """Clean and extract JSON from API response"""
        try:
            response_text = response_text.strip()
            
            # Remove markdown code blocks
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].strip()
            
            # Remove any text before first { and after last }
            start_idx = response_text.find('{')
            end_idx = response_text.rfind('}')
            
            if start_idx != -1 and end_idx != -1:
                response_text = response_text[start_idx:end_idx+1]
            
            return response_text
            
        except Exception as e:
            logger.error(f"Error cleaning JSON response: {str(e)}")
            raise ValueError("Invalid JSON format in API response")
    
    def _validate_json_structure(self, data: dict, required_fields: list) -> dict:
        """Validate JSON structure and provide defaults for missing fields"""
        validated_data = {}
        
        for field in required_fields:
            if field in data:
                validated_data[field] = data[field]
            else:
                # Provide sensible defaults
                if field in ['must_have_skills', 'good_to_have_skills', 'qualifications', 
                           'skills', 'education', 'certifications', 'projects', 'missing_skills',
                           'matching_skills', 'suggestions', 'strengths', 'improvement_areas']:
                    validated_data[field] = []
                elif field in ['title', 'name', 'experience_years', 'current_role', 'verdict']:
                    validated_data[field] = "Unknown"
                elif field == 'relevance_score':
                    validated_data[field] = 0
                else:
                    validated_data[field] = ""
        
        return validated_data
    
    @st.cache_data(ttl=3600)  # Cache for 1 hour
    def parse_job_description(_self, jd_text: str) -> Dict:
        """Parse job description to extract key components with caching"""
        try:
            validate_text_input(jd_text, min_length=50, max_length=20000)
            
            prompt = f"""
            Analyze the following job description and extract key information in JSON format:
            
            Job Description:
            {jd_text}
            
            Please extract and return ONLY a valid JSON object with these fields:
            {{
                "title": "extracted job title",
                "must_have_skills": ["skill1", "skill2", "skill3"],
                "good_to_have_skills": ["skill1", "skill2", "skill3"],
                "qualifications": ["qualification1", "qualification2"],
                "experience_required": "X years",
                "key_responsibilities": ["resp1", "resp2", "resp3"]
            }}
            
            Guidelines:
            - Extract 5-10 must-have skills (technical and soft skills)
            - Extract 3-7 good-to-have skills
            - Include education, certifications, and experience requirements
            - Focus on specific, actionable requirements
            - Return only valid JSON without additional text
            """
            
            response = _self.model.generate_content(prompt)
            if not response or not response.text:
                raise ValueError("Empty response from API")
                
            response_text = _self._clean_json_response(response.text)
            parsed_data = json.loads(response_text)
            
            required_fields = ['title', 'must_have_skills', 'good_to_have_skills', 
                             'qualifications', 'experience_required', 'key_responsibilities']
            
            validated_data = _self._validate_json_structure(parsed_data, required_fields)
            logger.info("Job description parsed successfully")
            return validated_data
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {str(e)}")
            st.error("Failed to parse API response. Please try again.")
            return _self._get_default_job_info()
        except Exception as e:
            logger.error(f"Error parsing job description: {str(e)}")
            st.error(f"Error analyzing job description: {str(e)}")
            return _self._get_default_job_info()
    
    def _get_default_job_info(self) -> Dict:
        """Return default job information structure"""
        return {
            "title": "Unknown",
            "must_have_skills": [],
            "good_to_have_skills": [],
            "qualifications": [],
            "experience_required": "Not specified",
            "key_responsibilities": []
        }
    
    @st.cache_data(ttl=1800)  # Cache for 30 minutes
    def extract_resume_info(_self, resume_text: str) -> Dict:
        """Extract information from resume with caching and validation"""
        try:
            validate_text_input(resume_text, min_length=100, max_length=50000)
            
            prompt = f"""
            Analyze the following resume and extract key information in JSON format:
            
            Resume:
            {resume_text}
            
            Please extract and return ONLY a valid JSON object with these fields:
            {{
                "name": "candidate name",
                "skills": ["skill1", "skill2", "skill3"],
                "experience_years": "X years or 0 if fresher",
                "education": ["degree1", "degree2"],
                "certifications": ["cert1", "cert2"],
                "projects": ["project1", "project2"],
                "current_role": "current job title or 'Fresher'",
                "email": "email@example.com",
                "phone": "phone number",
                "linkedin": "linkedin profile"
            }}
            
            Guidelines:
            - Extract all technical and soft skills mentioned
            - Calculate total years of experience accurately
            - Include all degrees, certifications, and training
            - List significant projects with brief descriptions
            - Return only valid JSON without additional text
            """
            
            response = _self.model.generate_content(prompt)
            if not response or not response.text:
                raise ValueError("Empty response from API")
                
            response_text = _self._clean_json_response(response.text)
            parsed_data = json.loads(response_text)
            
            required_fields = ['name', 'skills', 'experience_years', 'education', 
                             'certifications', 'projects', 'current_role', 'email', 'phone', 'linkedin']
            
            validated_data = _self._validate_json_structure(parsed_data, required_fields)
            logger.info("Resume information extracted successfully")
            return validated_data
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error in resume extraction: {str(e)}")
            st.error("Failed to parse resume analysis. Please try again.")
            return _self._get_default_resume_info()
        except Exception as e:
            logger.error(f"Error extracting resume info: {str(e)}")
            st.error(f"Error analyzing resume: {str(e)}")
            return _self._get_default_resume_info()
    
    def _get_default_resume_info(self) -> Dict:
        """Return default resume information structure"""
        return {
            "name": "Unknown",
            "skills": [],
            "experience_years": "0",
            "education": [],
            "certifications": [],
            "projects": [],
            "current_role": "Unknown",
            "email": "",
            "phone": "",
            "linkedin": ""
        }
    
    def calculate_relevance_score(self, resume_info: Dict, job_info: Dict) -> Tuple[int, Dict]:
        """Calculate relevance score and provide detailed analysis with enhanced validation"""
        try:
            if not resume_info or not job_info:
                raise ValueError("Resume or job information is missing")
            
            prompt = f"""
            Compare this resume with the job requirements and provide a comprehensive analysis:
            
            Resume Information:
            {json.dumps(resume_info, indent=2)}
            
            Job Requirements:
            {json.dumps(job_info, indent=2)}
            
            Please analyze and return ONLY a valid JSON object with:
            {{
                "relevance_score": 85,
                "verdict": "High/Medium/Low",
                "missing_skills": ["skill1", "skill2"],
                "matching_skills": ["skill1", "skill2"],
                "score_breakdown": {{
                    "skills_match": 40,
                    "experience_match": 25,
                    "education_match": 20,
                    "overall_fit": 15
                }},
                "suggestions": [
                    "Add certification in X",
                    "Gain experience in Y",
                    "Include projects related to Z"
                ],
                "strengths": ["strength1", "strength2"],
                "improvement_areas": ["area1", "area2"],
                "recommendations": [
                    "Specific action 1",
                    "Specific action 2"
                ]
            }}
            
            Scoring Guidelines:
            - Skills match (40 points): Technical and soft skills alignment
            - Experience match (25 points): Years and relevance of experience
            - Education match (20 points): Degree requirements and certifications
            - Overall fit (15 points): Cultural fit and growth potential
            
            Score Ranges:
            - 80-100: High suitability (Strong match, ready to hire)
            - 60-79: Medium suitability (Good potential with some gaps)
            - 0-59: Low suitability (Significant gaps, needs development)
            
            Return only valid JSON without additional text.
            """
            
            response = self.model.generate_content(prompt)
            if not response or not response.text:
                raise ValueError("Empty response from API")
                
            response_text = self._clean_json_response(response.text)
            analysis = json.loads(response_text)
            
            # Validate and clean the analysis
            required_fields = ['relevance_score', 'verdict', 'missing_skills', 'matching_skills',
                             'suggestions', 'strengths', 'improvement_areas']
            
            validated_analysis = self._validate_json_structure(analysis, required_fields)
            
            # Ensure score is within valid range
            score = validated_analysis.get('relevance_score', 0)
            if not isinstance(score, int) or score < 0 or score > 100:
                score = 0
                validated_analysis['relevance_score'] = 0
            
            # Validate verdict based on score
            if score >= 80:
                validated_analysis['verdict'] = 'High'
            elif score >= 60:
                validated_analysis['verdict'] = 'Medium'
            else:
                validated_analysis['verdict'] = 'Low'
            
            logger.info(f"Relevance score calculated: {score}")
            return score, validated_analysis
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error in score calculation: {str(e)}")
            st.error("Failed to parse scoring analysis. Please try again.")
            return 0, self._get_default_analysis()
        except Exception as e:
            logger.error(f"Error calculating relevance score: {str(e)}")
            st.error(f"Error calculating relevance score: {str(e)}")
            return 0, self._get_default_analysis()
    
    def _get_default_analysis(self) -> Dict:
        """Return default analysis structure"""
        return {
            "relevance_score": 0,
            "verdict": "Low",
            "missing_skills": [],
            "matching_skills": [],
            "suggestions": ["Unable to analyze resume - please try again"],
            "strengths": [],
            "improvement_areas": [],
            "score_breakdown": {
                "skills_match": 0,
                "experience_match": 0,
                "education_match": 0,
                "overall_fit": 0
            }
        }

class ResumeRelevanceSystem:
    """Main system class with enhanced database operations and caching"""
    
    def __init__(self, db_conn, analyzer: GeminiAnalyzer):
        self.db = db_conn
        self.analyzer = analyzer
        self.doc_processor = DocumentProcessor()
    
    def save_job_description(self, title: str, company: str, location: str, 
                           description: str, parsed_info: Dict) -> int:
        """Save job description to database with validation"""
        try:
            # Validate inputs
            validate_text_input(title, min_length=2, max_length=200)
            validate_text_input(description, min_length=50, max_length=20000)
            
            # Sanitize inputs
            title = title.strip()
            company = company.strip() if company else ""
            location = location.strip() if location else ""
            
            cursor = self.db.cursor()
            cursor.execute('''
                INSERT INTO job_descriptions 
                (title, company, location, description, must_have_skills, good_to_have_skills, qualifications)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                title, company, location, description,
                json.dumps(parsed_info.get('must_have_skills', [])),
                json.dumps(parsed_info.get('good_to_have_skills', [])),
                json.dumps(parsed_info.get('qualifications', []))
            ))
            
            self.db.commit()
            job_id = cursor.lastrowid
            logger.info(f"Job description saved with ID: {job_id}")
            return job_id
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Error saving job description: {str(e)}")
            raise ValueError(f"Failed to save job description: {str(e)}")
    
    def save_resume_evaluation(self, job_id: int, candidate_name: str, resume_text: str,
                              score: int, analysis: Dict, file_name: str = ""):
        """Save resume evaluation to database with enhanced data"""
        try:
            # Validate inputs
            if not candidate_name or not candidate_name.strip():
                candidate_name = "Unknown Candidate"
            
            validate_text_input(resume_text, min_length=50, max_length=100000)
            
            if not isinstance(score, int) or score < 0 or score > 100:
                score = 0
            
            # Sanitize filename
            safe_filename = sanitize_filename(file_name) if file_name else ""
            
            cursor = self.db.cursor()
            cursor.execute('''
                INSERT INTO resume_evaluations 
                (job_id, candidate_name, resume_text, relevance_score, missing_skills, 
                 matching_skills, verdict, suggestions, strengths, improvement_areas, file_name)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                job_id, candidate_name.strip(), resume_text, score,
                json.dumps(analysis.get('missing_skills', [])),
                json.dumps(analysis.get('matching_skills', [])),
                analysis.get('verdict', 'Low'),
                json.dumps(analysis.get('suggestions', [])),
                json.dumps(analysis.get('strengths', [])),
                json.dumps(analysis.get('improvement_areas', [])),
                safe_filename
            ))
            
            self.db.commit()
            evaluation_id = cursor.lastrowid
            logger.info(f"Resume evaluation saved with ID: {evaluation_id}")
            return evaluation_id
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Error saving resume evaluation: {str(e)}")
            raise ValueError(f"Failed to save resume evaluation: {str(e)}")
    
    @st.cache_data(ttl=600)  # Cache for 10 minutes
    def get_job_descriptions(_self) -> List[Dict]:
        """Get all job descriptions from database with caching"""
        try:
            cursor = _self.db.cursor()
            cursor.execute('''
                SELECT id, title, company, location, description, created_at, updated_at
                FROM job_descriptions 
                ORDER BY created_at DESC
            ''')
            
            jobs = []
            for row in cursor.fetchall():
                jobs.append({
                    'id': row[0],
                    'title': row[1],
                    'company': row[2],
                    'location': row[3],
                    'description': row[4],
                    'created_at': row[5],
                    'updated_at': row[6] if len(row) > 6 else row[5]
                })
            
            logger.info(f"Retrieved {len(jobs)} job descriptions")
            return jobs
            
        except Exception as e:
            logger.error(f"Error retrieving job descriptions: {str(e)}")
            st.error(f"Error loading job descriptions: {str(e)}")
            return []
    
    @st.cache_data(ttl=300)  # Cache for 5 minutes
    def get_evaluations_for_job(_self, job_id: int) -> List[Dict]:
        """Get all evaluations for a specific job with enhanced data"""
        try:
            cursor = _self.db.cursor()
            cursor.execute('''
                SELECT id, candidate_name, relevance_score, verdict, created_at, 
                       missing_skills, matching_skills, suggestions, file_name
                FROM resume_evaluations 
                WHERE job_id = ? 
                ORDER BY relevance_score DESC, created_at DESC
            ''', (job_id,))
            
            evaluations = []
            for row in cursor.fetchall():
                evaluations.append({
                    'id': row[0],
                    'candidate_name': row[1],
                    'relevance_score': row[2],
                    'verdict': row[3],
                    'created_at': row[4],
                    'missing_skills': json.loads(row[5]) if row[5] else [],
                    'matching_skills': json.loads(row[6]) if row[6] else [],
                    'suggestions': json.loads(row[7]) if row[7] else [],
                    'file_name': row[8] if len(row) > 8 else ""
                })
            
            logger.info(f"Retrieved {len(evaluations)} evaluations for job {job_id}")
            return evaluations
            
        except Exception as e:
            logger.error(f"Error retrieving evaluations: {str(e)}")
            st.error(f"Error loading evaluations: {str(e)}")
            return []
    
    def get_job_details(self, job_id: int) -> Optional[Dict]:
        """Get detailed job information by ID"""
        try:
            cursor = self.db.cursor()
            cursor.execute('''
                SELECT id, title, company, location, description, 
                       must_have_skills, good_to_have_skills, qualifications, created_at
                FROM job_descriptions 
                WHERE id = ?
            ''', (job_id,))
            
            row = cursor.fetchone()
            if row:
                return {
                    'id': row[0],
                    'title': row[1],
                    'company': row[2],
                    'location': row[3],
                    'description': row[4],
                    'must_have_skills': json.loads(row[5]) if row[5] else [],
                    'good_to_have_skills': json.loads(row[6]) if row[6] else [],
                    'qualifications': json.loads(row[7]) if row[7] else [],
                    'created_at': row[8]
                }
            return None
            
        except Exception as e:
            logger.error(f"Error retrieving job details: {str(e)}")
            return None
    
    def delete_job_description(self, job_id: int) -> bool:
        """Delete job description and associated evaluations"""
        try:
            cursor = self.db.cursor()
            
            # Delete evaluations first (due to foreign key constraint)
            cursor.execute('DELETE FROM resume_evaluations WHERE job_id = ?', (job_id,))
            
            # Delete job description
            cursor.execute('DELETE FROM job_descriptions WHERE id = ?', (job_id,))
            
            self.db.commit()
            
            if cursor.rowcount > 0:
                logger.info(f"Job description {job_id} deleted successfully")
                return True
            return False
            
        except Exception as e:
            self.db.rollback()
            logger.error(f"Error deleting job description: {str(e)}")
            return False
    
    def get_analytics_data(self) -> Dict:
        """Get comprehensive analytics data"""
        try:
            cursor = self.db.cursor()
            
            # Basic counts
            cursor.execute('SELECT COUNT(*) FROM job_descriptions')
            total_jobs = cursor.fetchone()[0]
            
            cursor.execute('SELECT COUNT(*) FROM resume_evaluations')
            total_evaluations = cursor.fetchone()[0]
            
            # Average score
            cursor.execute('SELECT AVG(relevance_score) FROM resume_evaluations')
            avg_score_result = cursor.fetchone()[0]
            avg_score = avg_score_result if avg_score_result else 0
            
            # Verdict distribution
            cursor.execute('SELECT verdict, COUNT(*) FROM resume_evaluations GROUP BY verdict')
            verdict_data = dict(cursor.fetchall())
            
            # Score distribution
            cursor.execute('SELECT relevance_score FROM resume_evaluations')
            scores = [row[0] for row in cursor.fetchall()]
            
            # Top performing candidates
            cursor.execute('''
                SELECT candidate_name, relevance_score, verdict, created_at
                FROM resume_evaluations 
                ORDER BY relevance_score DESC 
                LIMIT 10
            ''')
            top_candidates = [
                {
                    'name': row[0],
                    'score': row[1],
                    'verdict': row[2],
                    'date': row[3]
                }
                for row in cursor.fetchall()
            ]
            
            return {
                'total_jobs': total_jobs,
                'total_evaluations': total_evaluations,
                'avg_score': avg_score,
                'verdict_distribution': verdict_data,
                'all_scores': scores,
                'top_candidates': top_candidates
            }
            
        except Exception as e:
            logger.error(f"Error retrieving analytics data: {str(e)}")
            return {
                'total_jobs': 0,
                'total_evaluations': 0,
                'avg_score': 0,
                'verdict_distribution': {},
                'all_scores': [],
                'top_candidates': []
            }

def load_api_key():
    """Load API key from environment or user input with security measures"""
    # Try to load from environment first
    api_key = os.getenv('GEMINI_API_KEY')
    
    if api_key and len(api_key.strip()) > 10:
        return api_key.strip()
    
    # Use default API key if available
    if DEFAULT_GEMINI_API_KEY:
        return DEFAULT_GEMINI_API_KEY
    
    # If not in environment, ask user in sidebar
    st.sidebar.title("� Configuration")
    api_key = st.sidebar.text_input(
        "Enter Gemini API Key", 
        type="password",
        help="Get your API key from Google AI Studio",
        placeholder="Enter your API key here..."
    )
    
    return api_key.strip() if api_key else None

def create_settings_sidebar():
    """Create enhanced settings sidebar"""
    st.sidebar.markdown("---")
    st.sidebar.subheader("⚙️ Settings")
    
    # File upload settings
    max_file_size = st.sidebar.slider("Max File Size (MB)", 1, 50, 10)
    
    # Processing settings
    enable_caching = st.sidebar.checkbox("Enable Caching", value=True)
    
    # Display settings
    show_debug_info = st.sidebar.checkbox("Show Debug Info", value=False)
    
    # Export format
    export_format = st.sidebar.selectbox("Export Format", ["CSV", "Excel", "JSON"])
    
    return {
        'max_file_size': max_file_size,
        'enable_caching': enable_caching,
        'show_debug_info': show_debug_info,
        'export_format': export_format
    }

def display_system_info():
    """Display system information and health checks"""
    with st.expander("🔍 System Information"):
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.write("**Dependencies:**")
            for dep, available in DEPENDENCIES_AVAILABLE.items():
                status = "✅" if available else "❌"
                st.write(f"{status} {dep}")
        
        with col2:
            st.write("**Database:**")
            try:
                db_conn = init_database()
                cursor = db_conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
                tables = cursor.fetchall()
                st.write(f"✅ Tables: {len(tables)}")
            except Exception as e:
                st.write(f"❌ Database Error: {str(e)}")
        
        with col3:
            st.write("**Cache Status:**")
            cache_info = st.cache_data.get_stats()
            st.write(f"Cache hits: {len(cache_info)}")

def main():
    st.title("� Automated Resume Relevance Check System")
    st.markdown("### AI-Powered Resume Evaluation Platform")
    
    # Check dependencies first
    if not check_dependencies():
        st.stop()
    
    # Display system info in debug mode
    settings = create_settings_sidebar()
    if settings['show_debug_info']:
        display_system_info()
    
    # Initialize database
    try:
        db_conn = init_database()
    except Exception as e:
        st.error(f"Failed to initialize database: {str(e)}")
        st.stop()
    
    # Load API key
    api_key = load_api_key()
    
    if not api_key:
        st.warning("⚠️ API key configuration issue. Please check your setup.")
        st.info("💡 Get your free API key from: https://makersuite.google.com/app/apikey")
        
        with st.expander("📖 How to get your API key"):
            st.markdown("""
            1. Visit [Google AI Studio](https://makersuite.google.com/app/apikey)
            2. Sign in with your Google account
            3. Click "Create API Key"
            4. Copy the generated key
            5. Set it as environment variable or update the DEFAULT_GEMINI_API_KEY in the code
            
            **Security Note:** Your API key is used only for this session.
            """)
        st.stop()
    else:
        st.success("✅ API key loaded successfully!")
    
    # Initialize analyzer and system
    try:
        with st.spinner("Initializing AI analyzer..."):
            analyzer = GeminiAnalyzer(api_key)
            system = ResumeRelevanceSystem(db_conn, analyzer)
        st.success("✅ System initialized successfully!")
    except Exception as e:
        st.error(f"❌ System initialization failed: {str(e)}")
        st.stop()
    
    # Enhanced Navigation with improved styling
    st.markdown("---")
    
    # Quick stats overview
    analytics_data = system.get_analytics_data()
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📝 Job Descriptions", analytics_data['total_jobs'])
    with col2:
        st.metric("📄 Resume Evaluations", analytics_data['total_evaluations'])
    with col3:
        st.metric("📊 Average Score", f"{analytics_data['avg_score']:.1f}%")
    with col4:
        verdict_counts = analytics_data['verdict_distribution']
        high_count = verdict_counts.get('High', 0)
        st.metric("🎯 High Matches", high_count)
    
    # Main navigation tabs with enhanced styling
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📝 Job Management", 
        "📄 Resume Evaluation", 
        "📊 Dashboard", 
        "📈 Analytics",
        "⚙️ Settings"
    ])
    
    with tab1:
        st.header("📝 Job Description Management")
        
        # Add sub-tabs for better organization
        subtab1, subtab2 = st.tabs(["➕ Add New Job", "📋 Manage Existing Jobs"])
        
        with subtab1:
            st.subheader("Add New Job Description")
            
            with st.form("job_form", clear_on_submit=True):
                col1, col2 = st.columns(2)
                
                with col1:
                    job_title = st.text_input(
                        "Job Title *", 
                        placeholder="e.g., Senior Python Developer",
                        help="Enter the exact job title"
                    )
                    company_name = st.text_input(
                        "Company Name", 
                        placeholder="e.g., Tech Corp",
                        help="Company or organization name"
                    )
                
                with col2:
                    location = st.text_input(
                        "Location", 
                        placeholder="e.g., Hyderabad, India / Remote",
                        help="Job location or remote work option"
                    )
                    job_type = st.selectbox(
                        "Job Type",
                        ["Full-time", "Part-time", "Contract", "Internship", "Remote"]
                    )
                
                jd_text = st.text_area(
                    "Job Description *", 
                    height=300,
                    placeholder="Paste the complete job description here...",
                    help="Include responsibilities, requirements, skills, and qualifications"
                )
                
                submit_button = st.form_submit_button("🚀 Parse & Save Job Description", type="primary")
                
                if submit_button:
                    if job_title and jd_text:
                        try:
                            with st.spinner("🔍 Analyzing job description..."):
                                parsed_info = analyzer.parse_job_description(jd_text)
                                job_id = system.save_job_description(
                                    job_title, company_name, location, jd_text, parsed_info
                                )
                            
                            st.success(f"✅ Job description saved successfully! (ID: {job_id})")
                            
                            # Display parsed information with better formatting
                            st.subheader("🎯 Extracted Information")
                            
                            col_a, col_b, col_c = st.columns(3)
                            
                            with col_a:
                                st.write("**🔧 Must-Have Skills:**")
                                if parsed_info.get('must_have_skills'):
                                    for skill in parsed_info['must_have_skills']:
                                        st.write(f"• {skill}")
                                else:
                                    st.write("None specified")
                            
                            with col_b:
                                st.write("**💡 Good-to-Have Skills:**")
                                if parsed_info.get('good_to_have_skills'):
                                    for skill in parsed_info['good_to_have_skills']:
                                        st.write(f"• {skill}")
                                else:
                                    st.write("None specified")
                            
                            with col_c:
                                st.write("**🎓 Qualifications:**")
                                if parsed_info.get('qualifications'):
                                    for qual in parsed_info['qualifications']:
                                        st.write(f"• {qual}")
                                else:
                                    st.write("None specified")
                            
                            # Clear cache to refresh job list
                            st.cache_data.clear()
                            
                        except Exception as e:
                            st.error(f"❌ Error processing job description: {str(e)}")
                    else:
                        st.error("⚠️ Please fill in the required fields (Job Title and Description).")
        
        with subtab2:
            st.subheader("📋 Existing Job Descriptions")
            
            jobs = system.get_job_descriptions()
            
            if jobs:
                # Add search and filter options
                col1, col2 = st.columns([3, 1])
                
                with col1:
                    search_term = st.text_input("🔍 Search jobs", placeholder="Search by title, company, or location...")
                
                with col2:
                    sort_option = st.selectbox("Sort by", ["Date (Newest)", "Date (Oldest)", "Title A-Z", "Company A-Z"])
                
                # Filter and sort jobs
                filtered_jobs = jobs
                if search_term:
                    filtered_jobs = [
                        job for job in jobs
                        if search_term.lower() in job['title'].lower() 
                        or search_term.lower() in job['company'].lower()
                        or search_term.lower() in job['location'].lower()
                    ]
                
                # Display jobs in cards
                for job in filtered_jobs[:10]:  # Show first 10 jobs
                    with st.expander(f"📋 {job['title']} - {job['company']}", expanded=False):
                        col1, col2, col3 = st.columns([2, 1, 1])
                        
                        with col1:
                            st.write(f"**📍 Location:** {job['location']}")
                            st.write(f"**📅 Posted:** {job['created_at']}")
                            st.write(f"**🆔 ID:** {job['id']}")
                        
                        with col2:
                            evaluations = system.get_evaluations_for_job(job['id'])
                            st.metric("Applications", len(evaluations))
                        
                        with col3:
                            if st.button(f"🗑️ Delete", key=f"delete_{job['id']}", type="secondary"):
                                if system.delete_job_description(job['id']):
                                    st.success("Job deleted successfully!")
                                    st.cache_data.clear()
                                    st.rerun()
                                else:
                                    st.error("Failed to delete job")
                
                if len(filtered_jobs) > 10:
                    st.info(f"Showing 10 of {len(filtered_jobs)} jobs. Use search to narrow results.")
            else:
                st.info("🌟 No job descriptions found. Add your first job to get started!")
    
    with tab2:
        st.header("📄 Resume Evaluation Center")
        
        # Job selection with better UI
        jobs = system.get_job_descriptions()
        if not jobs:
            st.warning("⚠️ Please add at least one job description before evaluating resumes.")
            return
        
        st.subheader("🎯 Select Target Position")
        job_options = {
            f"{job['title']} - {job['company']} {'(' + job['location'] + ')' if job['location'] else ''} [ID: {job['id']}]": job['id'] 
            for job in jobs
        }
        
        selected_job_display = st.selectbox(
            "Choose the position to evaluate against:",
            options=list(job_options.keys()),
            help="Select the job description to compare resumes against"
        )
        selected_job_id = job_options[selected_job_display]
        
        # Display job summary
        selected_job = system.get_job_details(selected_job_id)
        if selected_job:
            with st.expander("📋 Job Requirements Summary", expanded=True):
                col1, col2 = st.columns(2)
                with col1:
                    st.write("**🔧 Must-Have Skills:**")
                    for skill in selected_job['must_have_skills'][:5]:
                        st.write(f"• {skill}")
                    if len(selected_job['must_have_skills']) > 5:
                        st.write(f"• ... and {len(selected_job['must_have_skills']) - 5} more")
                
                with col2:
                    st.write("**🎓 Key Qualifications:**")
                    for qual in selected_job['qualifications'][:3]:
                        st.write(f"• {qual}")
                    if len(selected_job['qualifications']) > 3:
                        st.write(f"• ... and {len(selected_job['qualifications']) - 3} more")
        
        st.markdown("---")
        
        # Enhanced file upload section
        upload_tab1, upload_tab2 = st.tabs(["📄 Single Resume", "📁 Batch Upload"])
        
        with upload_tab1:
            st.subheader("📄 Single Resume Evaluation")
            
            uploaded_file = st.file_uploader(
                "Upload Resume", 
                type=['pdf', 'docx'],
                accept_multiple_files=False,
                help=f"Supported formats: PDF, DOCX (Max size: {settings['max_file_size']}MB)"
            )
            
            if uploaded_file:
                # Display file info
                st.info(f"📁 File: {uploaded_file.name} ({uploaded_file.size / 1024:.1f} KB)")
                
                if st.button("🔍 Evaluate Resume", type="primary", use_container_width=True):
                    try:
                        with st.spinner("🔄 Processing resume..."):
                            # Progress tracking
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            # Step 1: Extract text
                            status_text.text("📖 Extracting text from document...")
                            progress_bar.progress(25)
                            resume_text = DocumentProcessor.extract_text_from_file(uploaded_file)
                            
                            if not resume_text:
                                st.error("❌ Could not extract text from the resume. Please check the file format.")
                                return
                            
                            # Step 2: Extract resume info
                            status_text.text("🧠 Analyzing resume content...")
                            progress_bar.progress(50)
                            resume_info = analyzer.extract_resume_info(resume_text)
                            
                            # Step 3: Get job info
                            status_text.text("📋 Loading job requirements...")
                            progress_bar.progress(75)
                            job_info = {
                                'title': selected_job['title'],
                                'must_have_skills': selected_job['must_have_skills'],
                                'good_to_have_skills': selected_job['good_to_have_skills'],
                                'qualifications': selected_job['qualifications']
                            }
                            
                            # Step 4: Calculate score
                            status_text.text("⚡ Calculating relevance score...")
                            progress_bar.progress(90)
                            score, analysis = analyzer.calculate_relevance_score(resume_info, job_info)
                            
                            # Step 5: Save results
                            status_text.text("💾 Saving evaluation results...")
                            progress_bar.progress(100)
                            system.save_resume_evaluation(
                                selected_job_id, 
                                resume_info.get('name', 'Unknown'),
                                resume_text, 
                                score, 
                                analysis, 
                                uploaded_file.name
                            )
                            
                            # Clear progress indicators
                            progress_bar.empty()
                            status_text.empty()
                        
                        st.success("✅ Resume evaluation completed successfully!")
                        
                        # Enhanced results display
                        st.markdown("---")
                        st.subheader("📊 Evaluation Results")
                        
                        # Score display with color coding
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            score_color = "🟢" if score >= 80 else "🟡" if score >= 60 else "🔴"
                            st.metric("Relevance Score", f"{score}%", delta=None)
                            st.markdown(f"<h3 style='text-align: center;'>{score_color}</h3>", unsafe_allow_html=True)
                        
                        with col2:
                            verdict_colors = {"High": "🟢", "Medium": "🟡", "Low": "🔴"}
                            verdict_color = verdict_colors.get(analysis['verdict'], '⚪')
                            st.metric("Verdict", f"{verdict_color} {analysis['verdict']}")
                        
                        with col3:
                            st.metric("Candidate", resume_info.get('name', 'Unknown'))
                        
                        with col4:
                            matching_skills = len(analysis.get('matching_skills', []))
                            total_required = len(selected_job['must_have_skills'])
                            if total_required > 0:
                                skill_match_percent = (matching_skills / total_required) * 100
                                st.metric("Skill Match", f"{skill_match_percent:.0f}%")
                            else:
                                st.metric("Skills Found", matching_skills)
                        
                        # Detailed analysis in tabs
                        detail_tab1, detail_tab2, detail_tab3, detail_tab4 = st.tabs([
                            "✅ Strengths", "❌ Gaps", "💡 Suggestions", "📈 Breakdown"
                        ])
                        
                        with detail_tab1:
                            col1, col2 = st.columns(2)
                            with col1:
                                st.write("**🎯 Matching Skills:**")
                                for skill in analysis.get('matching_skills', []):
                                    st.write(f"✅ {skill}")
                            
                            with col2:
                                st.write("**💪 Key Strengths:**")
                                for strength in analysis.get('strengths', []):
                                    st.write(f"🌟 {strength}")
                        
                        with detail_tab2:
                            col1, col2 = st.columns(2)
                            with col1:
                                st.write("**❌ Missing Skills:**")
                                for skill in analysis.get('missing_skills', []):
                                    st.write(f"⚠️ {skill}")
                            
                            with col2:
                                st.write("**📈 Improvement Areas:**")
                                for area in analysis.get('improvement_areas', []):
                                    st.write(f"🔧 {area}")
                        
                        with detail_tab3:
                            st.write("**💡 Recommendations:**")
                            for i, suggestion in enumerate(analysis.get('suggestions', []), 1):
                                st.write(f"{i}. {suggestion}")
                        
                        with detail_tab4:
                            if 'score_breakdown' in analysis:
                                breakdown = analysis['score_breakdown']
                                
                                # Create a simple bar chart
                                breakdown_df = pd.DataFrame([
                                    {'Category': 'Skills Match', 'Score': breakdown.get('skills_match', 0)},
                                    {'Category': 'Experience', 'Score': breakdown.get('experience_match', 0)},
                                    {'Category': 'Education', 'Score': breakdown.get('education_match', 0)},
                                    {'Category': 'Overall Fit', 'Score': breakdown.get('overall_fit', 0)}
                                ])
                                
                                st.bar_chart(breakdown_df.set_index('Category'))
                                
                                # Display detailed breakdown
                                for category, score in breakdown.items():
                                    category_name = category.replace('_', ' ').title()
                                    st.write(f"**{category_name}:** {score}/40" if 'skills' in category else f"**{category_name}:** {score}")
                        
                        # Clear cache to refresh data
                        st.cache_data.clear()
                        
                    except Exception as e:
                        st.error(f"❌ Error during evaluation: {str(e)}")
                        logger.error(f"Evaluation error: {str(e)}")
        
        with upload_tab2:
            st.subheader("📁 Batch Resume Processing")
            
            uploaded_files = st.file_uploader(
                "Upload Multiple Resumes", 
                type=['pdf', 'docx'],
                accept_multiple_files=True,
                help=f"Upload multiple resume files for batch processing (Max {settings['max_file_size']}MB each)"
            )
            
            if uploaded_files:
                st.info(f"📁 Selected {len(uploaded_files)} files for processing")
                
                # Display file list
                with st.expander("📋 File List"):
                    for i, file in enumerate(uploaded_files, 1):
                        st.write(f"{i}. {file.name} ({file.size / 1024:.1f} KB)")
                
                if st.button("🚀 Process All Resumes", type="primary", use_container_width=True):
                    results = []
                    
                    # Create containers for progress tracking
                    progress_container = st.container()
                    results_container = st.container()
                    
                    with progress_container:
                        overall_progress = st.progress(0)
                        current_file_text = st.empty()
                        file_progress = st.progress(0)
                    
                    # Get job info once
                    job_info = {
                        'title': selected_job['title'],
                        'must_have_skills': selected_job['must_have_skills'],
                        'good_to_have_skills': selected_job['good_to_have_skills'],
                        'qualifications': selected_job['qualifications']
                    }
                    
                    for i, uploaded_file in enumerate(uploaded_files):
                        try:
                            current_file_text.text(f"Processing: {uploaded_file.name}")
                            
                            # Extract text
                            file_progress.progress(25)
                            resume_text = DocumentProcessor.extract_text_from_file(uploaded_file)
                            
                            if resume_text:
                                # Extract resume info
                                file_progress.progress(50)
                                resume_info = analyzer.extract_resume_info(resume_text)
                                
                                # Calculate score
                                file_progress.progress(75)
                                score, analysis = analyzer.calculate_relevance_score(resume_info, job_info)
                                
                                # Save evaluation
                                file_progress.progress(90)
                                system.save_resume_evaluation(
                                    selected_job_id,
                                    resume_info.get('name', f'Candidate_{i+1}'),
                                    resume_text,
                                    score,
                                    analysis,
                                    uploaded_file.name
                                )
                                
                                results.append({
                                    'file_name': uploaded_file.name,
                                    'candidate_name': resume_info.get('name', f'Candidate_{i+1}'),
                                    'score': score,
                                    'verdict': analysis['verdict'],
                                    'matching_skills': len(analysis.get('matching_skills', [])),
                                    'missing_skills': len(analysis.get('missing_skills', []))
                                })
                            else:
                                results.append({
                                    'file_name': uploaded_file.name,
                                    'candidate_name': 'Error',
                                    'score': 0,
                                    'verdict': 'Error',
                                    'matching_skills': 0,
                                    'missing_skills': 0
                                })
                            
                            file_progress.progress(100)
                            overall_progress.progress((i + 1) / len(uploaded_files))
                            
                        except Exception as e:
                            logger.error(f"Error processing {uploaded_file.name}: {str(e)}")
                            results.append({
                                'file_name': uploaded_file.name,
                                'candidate_name': 'Error',
                                'score': 0,
                                'verdict': 'Error',
                                'matching_skills': 0,
                                'missing_skills': 0
                            })
                    
                    # Clear progress indicators
                    progress_container.empty()
                    
                    # Display results
                    with results_container:
                        st.success(f"✅ Processed {len(results)} resumes!")
                        
                        # Summary statistics
                        successful_results = [r for r in results if r['verdict'] != 'Error']
                        if successful_results:
                            col1, col2, col3, col4 = st.columns(4)
                            
                            with col1:
                                avg_score = sum(r['score'] for r in successful_results) / len(successful_results)
                                st.metric("Average Score", f"{avg_score:.1f}%")
                            
                            with col2:
                                high_count = len([r for r in successful_results if r['verdict'] == 'High'])
                                st.metric("High Matches", high_count)
                            
                            with col3:
                                medium_count = len([r for r in successful_results if r['verdict'] == 'Medium'])
                                st.metric("Medium Matches", medium_count)
                            
                            with col4:
                                error_count = len([r for r in results if r['verdict'] == 'Error'])
                                st.metric("Processing Errors", error_count)
                        
                        # Results table with sorting
                        df = pd.DataFrame(results)
                        df = df.sort_values('score', ascending=False)
                        
                        # Color code the dataframe
                        def highlight_verdict(val):
                            colors = {
                                'High': 'background-color: #d4edda',
                                'Medium': 'background-color: #fff3cd', 
                                'Low': 'background-color: #f8d7da',
                                'Error': 'background-color: #f5c6cb'
                            }
                            return colors.get(val, '')
                        
                        styled_df = df.style.applymap(highlight_verdict, subset=['verdict'])
                        st.dataframe(styled_df, use_container_width=True)
                        
                        # Export options
                        if successful_results:
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                # Export all results
                                csv_all = df.to_csv(index=False)
                                st.download_button(
                                    label="📥 Download All Results (CSV)",
                                    data=csv_all,
                                    file_name=f"batch_evaluation_all_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                    mime="text/csv"
                                )
                            
                            with col2:
                                # Export only high/medium matches
                                shortlist = df[df['verdict'].isin(['High', 'Medium'])]
                                if not shortlist.empty:
                                    csv_shortlist = shortlist.to_csv(index=False)
                                    st.download_button(
                                        label="🎯 Download Shortlist (CSV)",
                                        data=csv_shortlist,
                                        file_name=f"shortlist_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                                        mime="text/csv"
                                    )
                        
                        # Clear cache
                        st.cache_data.clear()
    
    with tab3:
        st.header("📊 Evaluation Dashboard")
        
        # Job selection for dashboard
        jobs = system.get_job_descriptions()
        if jobs:
            dashboard_col1, dashboard_col2 = st.columns([3, 1])
            
            with dashboard_col1:
                job_options = {f"{job['title']} - {job['company']}": job['id'] for job in jobs}
                selected_job_display = st.selectbox(
                    "Select Job for Dashboard", 
                    options=list(job_options.keys()), 
                    key="dashboard_job"
                )
                selected_job_id = job_options[selected_job_display]
            
            with dashboard_col2:
                auto_refresh = st.checkbox("Auto Refresh", value=False)
                if auto_refresh:
                    st.rerun()
            
            evaluations = system.get_evaluations_for_job(selected_job_id)
            
            if evaluations:
                # Enhanced summary metrics
                col1, col2, col3, col4, col5 = st.columns(5)
                
                with col1:
                    st.metric("Total Applications", len(evaluations))
                
                with col2:
                    high_count = len([e for e in evaluations if e['verdict'] == 'High'])
                    st.metric("High Suitability", high_count, delta=f"{(high_count/len(evaluations)*100):.1f}%")
                
                with col3:
                    medium_count = len([e for e in evaluations if e['verdict'] == 'Medium'])
                    st.metric("Medium Suitability", medium_count, delta=f"{(medium_count/len(evaluations)*100):.1f}%")
                
                with col4:
                    avg_score = sum([e['relevance_score'] for e in evaluations]) / len(evaluations)
                    st.metric("Average Score", f"{avg_score:.1f}%")
                
                with col5:
                    top_score = max([e['relevance_score'] for e in evaluations])
                    st.metric("Top Score", f"{top_score}%")
                
                # Interactive charts
                st.subheader("📈 Score Distribution")
                
                # Create score distribution chart
                scores = [e['relevance_score'] for e in evaluations]
                score_ranges = {
                    '90-100': len([s for s in scores if 90 <= s <= 100]),
                    '80-89': len([s for s in scores if 80 <= s < 90]),
                    '70-79': len([s for s in scores if 70 <= s < 80]),
                    '60-69': len([s for s in scores if 60 <= s < 70]),
                    '50-59': len([s for s in scores if 50 <= s < 60]),
                    'Below 50': len([s for s in scores if s < 50])
                }
                
                chart_col1, chart_col2 = st.columns(2)
                
                with chart_col1:
                    # Score distribution bar chart
                    range_df = pd.DataFrame(list(score_ranges.items()), columns=['Score Range', 'Count'])
                    st.bar_chart(range_df.set_index('Score Range'))
                
                with chart_col2:
                    # Verdict pie chart data
                    verdict_counts = {}
                    for evaluation in evaluations:
                        verdict = evaluation['verdict']
                        verdict_counts[verdict] = verdict_counts.get(verdict, 0) + 1
                    
                    verdict_df = pd.DataFrame(list(verdict_counts.items()), columns=['Verdict', 'Count'])
                    st.write("**Verdict Distribution**")
                    for _, row in verdict_df.iterrows():
                        percentage = (row['Count'] / len(evaluations)) * 100
                        st.write(f"**{row['Verdict']}:** {row['Count']} ({percentage:.1f}%)")
                
                # Top candidates table
                st.subheader("🏆 Top Candidates")
                
                # Filter and display options
                filter_col1, filter_col2, filter_col3 = st.columns(3)
                
                with filter_col1:
                    min_score = st.slider("Minimum Score", 0, 100, 60)
                
                with filter_col2:
                    verdict_filter = st.multiselect(
                        "Filter by Verdict", 
                        ['High', 'Medium', 'Low'],
                        default=['High', 'Medium']
                    )
                
                with filter_col3:
                    show_count = st.selectbox("Show Top", [10, 20, 50, 100], index=0)
                
                # Apply filters
                filtered_evaluations = [
                    e for e in evaluations 
                    if e['relevance_score'] >= min_score and e['verdict'] in verdict_filter
                ]
                
                # Display filtered results
                if filtered_evaluations:
                    display_evaluations = filtered_evaluations[:show_count]
                    
                    # Create enhanced table
                    table_data = []
                    for i, eval in enumerate(display_evaluations, 1):
                        table_data.append({
                            'Rank': i,
                            'Candidate': eval['candidate_name'],
                            'Score': eval['relevance_score'],
                            'Verdict': eval['verdict'],
                            'File': eval.get('file_name', 'N/A'),
                            'Date': eval['created_at'],
                            'Matching Skills': len(eval.get('matching_skills', [])),
                            'Missing Skills': len(eval.get('missing_skills', []))
                        })
                    
                    table_df = pd.DataFrame(table_data)
                    
                    # Style the table
                    def highlight_rows(row):
                        if row['Verdict'] == 'High':
                            return ['background-color: #d4edda'] * len(row)
                        elif row['Verdict'] == 'Medium':
                            return ['background-color: #fff3cd'] * len(row)
                        else:
                            return ['background-color: #f8d7da'] * len(row)
                    
                    styled_table = table_df.style.apply(highlight_rows, axis=1)
                    st.dataframe(styled_table, use_container_width=True)
                    
                    # Download shortlisted candidates
                    if len(display_evaluations) > 0:
                        csv_data = table_df.to_csv(index=False)
                        st.download_button(
                            label="📥 Download Results (CSV)",
                            data=csv_data,
                            file_name=f"candidates_{selected_job_display.replace(' ', '_')}.csv",
                            mime="text/csv"
                        )
                
                else:
                    st.info("No candidates match the selected filters.")
            
            else:
                st.info("📝 No evaluations found for this job position.")
        
        else:
            st.warning("📋 No job descriptions available.")
    
    with tab4:
        st.header("📈 Advanced Analytics")
        
        # Get comprehensive analytics data
        analytics_data = system.get_analytics_data()
        
        if analytics_data['total_evaluations'] > 0:
            # System overview
            overview_col1, overview_col2, overview_col3, overview_col4 = st.columns(4)
            
            with overview_col1:
                st.metric("Total Jobs", analytics_data['total_jobs'])
            
            with overview_col2:
                st.metric("Total Applications", analytics_data['total_evaluations'])
            
            with overview_col3:
                st.metric("Average Score", f"{analytics_data['avg_score']:.1f}%")
            
            with overview_col4:
                high_quality = analytics_data['verdict_distribution'].get('High', 0)
                st.metric("High Quality", high_quality)
            
            # Score distribution analysis
            st.subheader("📊 Score Analysis")
            
            scores = analytics_data['all_scores']
            if scores:
                score_ranges = {
                    '80-100 (High)': len([s for s in scores if s >= 80]),
                    '60-79 (Medium)': len([s for s in scores if 60 <= s < 80]),
                    '0-59 (Low)': len([s for s in scores if s < 60])
                }
                
                range_df = pd.DataFrame(list(score_ranges.items()), columns=['Score Range', 'Count'])
                st.bar_chart(range_df.set_index('Score Range'))
            
            # Top performers
            st.subheader("🏆 Top Performers")
            
            top_candidates = analytics_data['top_candidates']
            if top_candidates:
                top_df = pd.DataFrame(top_candidates)
                st.dataframe(
                    top_df[['name', 'score', 'verdict']].rename(columns={
                        'name': 'Candidate',
                        'score': 'Score',
                        'verdict': 'Verdict'
                    }),
                    use_container_width=True
                )
        
        else:
            st.info("📊 No analytics data available yet.")
    
    with tab5:
        st.header("⚙️ Settings")
        
        # API Configuration
        st.subheader("🔑 API Settings")
        current_key_status = "✅ Connected" if api_key else "❌ Not Connected"
        st.write(f"**Status:** {current_key_status}")
        
        # Cache Management
        st.subheader("🧹 Cache Management")
        if st.button("Clear Cache"):
            st.cache_data.clear()
            st.success("Cache cleared!")
        
        # Data Export
        st.subheader("📥 Data Export")
        
        export_col1, export_col2 = st.columns(2)
        
        with export_col1:
            if st.button("Export Jobs"):
                jobs = system.get_job_descriptions()
                if jobs:
                    jobs_df = pd.DataFrame(jobs)
                    csv_data = jobs_df.to_csv(index=False)
                    
                    st.download_button(
                        label="Download Jobs CSV",
                        data=csv_data,
                        file_name=f"jobs_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv"
                    )
        
        with export_col2:
            if st.button("Export Analytics"):
                csv_data = pd.DataFrame([analytics_data]).to_csv(index=False)
                
                st.download_button(
                    label="Download Analytics CSV",
                    data=csv_data,
                    file_name=f"analytics_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv"
                )

if __name__ == "__main__":
    main()